#-*- encoding=utf-8 -*-
import os
import sys
import subprocess
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import datetime
#from datetime import datetime
import pymysql
from PyPDF2 import PdfFileReader as pdfReader
from PyPDF2 import PdfFileWriter as pdfWriter
import time
import traceback
import re

def pdf_add_pswd(input_file, output_file, pswd):
    in_file = open(input_file, "rb")
    in_pdf = pdfReader(in_file)
    out_pdf = pdfWriter()
    p = in_pdf.getPage(0)
    out_pdf.addPage(p)
    out_pdf.encrypt(pswd)
    out_file = open(output_file, "wb")
    out_pdf.write(out_file)
    in_file.close()
    out_file.close()

#連線mysql
db_settings = {
                "host": "localhost",
                "port": 3306,
                "user": "libo_user",
                "password": "xxx",
                "db": "libodb",
            }
# db_settings = {
#                 "host": "192.168.2.115",
#                 "port": 3306,
#                 "user": "root",
#                 "password": "password",
#                 "db": "libodb",
#             }
try:
    # 建立Connection物件
    conn = pymysql.connect(**db_settings)
except Exception as ex:
    print(ex)
cursor = conn.cursor()

testid = sys.argv[1]
cmd = ''

# sample_id=str(sample_name)
# sample_name_full =re.match(r"PC|NTC|^Q[0-9]{9}|^QL[0-9]{8}|^F[0-9]{9}", sample_id)
# if sample_name_full!=None:
#     sample_id=sample_name_full[0]
# else:
#     if len(sample_id)!=15:
#         sample_id="0"*(15-len(sample_id))+sample_id
#     else:
#         sample_id=sample_id

if testid[0] == 'Q' or (re.match(r"[0-9]",testid[0])!=None):
    cmd = ("SELECT uuid,cname,fname,lname,sex,dob,mobile,userid,residentpermit,passportid,uemail,tdat,type,rdat,ftest,pcrtest,vuser2,hicardno,sendname,mtpid " 
            "FROM covid_trans WHERE sampleid2 = '" + testid + "'")
elif testid[0] == 'F':
    cmd = ("SELECT uuid,cname,fname,lname,sex,dob,mobile,userid,residentpermit,passportid,uemail,tdat,type,rdat,ftest,pcrtest,vuser2,hicardno,sendname,mtpid " 
            "FROM covid_trans WHERE sampleid1 = '" + testid + "'")
else:
    print("<script>alert('快篩ID或PCR ID格式錯誤');</script>")
    print("<script>alert('快篩ID為F開頭，PCR ID為Q開頭');</script>")

if cmd != '':
    cursor.execute(cmd)
    uuid,cname,fname,lname,sex,dob,mobile,userid,residentpermit,passportid,uemail,tdat,app_type,rdat,ftest,pcrtest,vuser2,hicardno,sendname,mtpid = cursor.fetchone()[:]
    dob = str(dob)
    tdat = str(tdat)
    rdat = str(rdat)
    fname = fname.upper()
    lname = lname.upper()
    if userid.strip() != '' and userid.strip() != 'NA':
        pwd = userid.strip()
    # elif residentpermit == 'Y':
    #     pwd = hicardno.strip()
    elif passportid.strip() != '' and passportid.strip() != 'NA':
        pwd = passportid.strip()
    else:
        pwd = dob.strip().replace('-', '').replace('/', '')
    if testid[0] == 'Q' or (re.match(r"[0-9]",testid[0])!=None):
        #打開模板
        doc = docx.Document('/var/www/html/templates/covid_templates_v31.docx')
        doc.styles['Normal'].font.name = u'微軟正黑體'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')
        
        Tables = doc.tables
        
        #table
        table0 = Tables[0]
        table1 = Tables[1]
        
        #中英姓名
        ename = fname + ' ' + lname
        if cname.strip() != '' and ename.strip() != '':
            run = table0.cell(2,1).paragraphs[0].add_run(cname + ' / ' + ename.strip())
        elif ename.strip() == '':
            run = table0.cell(2,1).paragraphs[0].add_run(cname)
        elif cname.strip() == '':
            run = table0.cell(2,1).paragraphs[0].add_run(ename)
        else:
            pass
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(2,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #生日
        run = table0.cell(2,3).paragraphs[0].add_run(dob.replace('-', '.'))
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(2,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #性別
        run = table0.cell(3,1).paragraphs[0].add_run(sex)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(3,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #手機
        run = table0.cell(3,3).paragraphs[0].add_run(mobile)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(3,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #身分證或居留證
        if userid.strip() != '':
            idcardno = userid.strip()
            run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
            run.font.size = Pt(9)
            run.font.bold = True
            table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.font.bold = True
            table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        elif residentpermit.strip() == 'Y':
            idcardno = hicardno.strip()
            run = table0.cell(4,0).paragraphs[0].add_run("居留證證號")
            run.font.size = Pt(9)
            run.font.bold = True
            table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = table0.cell(4,0).paragraphs[1].add_run("(ARC No.)")
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.font.bold = True
            table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            idcardno = 'NA'
            run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
            run.font.size = Pt(9)
            run.font.bold = True
            table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.font.bold = True
            table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        run = table0.cell(4,1).paragraphs[0].add_run(idcardno)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(4,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #護照
        if passportid.strip() == '':
            passportid = 'NA'
        else:
            pass
        run = table0.cell(4,3).paragraphs[0].add_run(passportid)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(4,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #e-mail
        if uemail.strip() == '':
            uemail = 'NA'
        else:
            pass
        run = table0.cell(5,1).paragraphs[0].add_run(uemail)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(5,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        #鼻咽
        # if xmrpturgency == '1':
        #     run = table0.cell(5,3).paragraphs[0].add_run("鼻咽拭子(Nasopharyngeal Swab)")
        #     run.font.size = Pt(9)
        #     run.font.name = 'Times New Roman'
        #     run.font.bold = True
        #     table0.cell(5,3).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        #     table0.cell(5,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # else:
        #     run = table0.cell(5,3).paragraphs[0].add_run("咽喉拭子 (Oropharyngeal Swab)")
        #     run.font.size = Pt(9)
        #     run.font.name = 'Times New Roman'
        #     run.font.bold = True
        #     table0.cell(5,3).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        #     table0.cell(5,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # run = table0.cell(5,3).paragraphs[0].add_run("咽喉拭子 (Oropharyngeal Swab)")
        # run.font.size = Pt(9)
        # run.font.name = 'Times New Roman'
        # run.font.bold = True
        # table0.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # table0.cell(5,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #採樣日期
        tdate = tdat.split(' ')[0]
        tdatime = tdat[:-3]
        run = table0.cell(6,1).paragraphs[0].add_run(tdatime)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(6,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #報告類型
        #if app_type.strip() == '1':
        #    app_type = '個人'
        #elif app_type.strip() == '2':
        #    cmd = "SELECT companyname from company where companyid='" + sendname + "'"
        #    cursor.execute(cmd)
        #    company = cursor.fetchone()
        #    if company == 'None' or company is None:
        #        app_type = '團體'
        #    else:
        #        app_type = company[0]
                
        #run = table0.cell(6,3).paragraphs[0].add_run(app_type)
        #run.font.size = Pt(9)
        #run.font.name = 'Times New Roman'
        #run.font.bold = True
        #table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        #table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        #台胞證
        if mtpid.strip() == '':
            mtpid = 'NA'
        else:
            pass
        run = table0.cell(6,3).paragraphs[0].add_run(mtpid)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        
        #檢測日期 ((報告時間+報到時間)/2)
        inspectime = datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(minutes = 80)
        if int(inspectime.strftime('%Y%m%d%H%M%S')) > int(tdat.replace('-', '').replace(':', '').replace(' ', '')):
            inspectime = inspectime.strftime('%Y-%m-%d %H:%M')
        else:
            inspectime = (datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(seconds = round(int((datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.datetime.strptime(tdat, '%Y-%m-%d %H:%M:%S')).seconds) / 2, 0))).strftime('%Y-%m-%d %H:%M')
        run = table0.cell(7,1).paragraphs[0].add_run(inspectime)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(7,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #檢測項目、檢測類型、方法、結果、檔案名
        testtype = "新冠肺炎病毒核酸檢測\n(SARS-CoV-2 Virus PCR Testing)"
        method = '即時反轉錄聚合酶連鎖反應 (Real-time RT-PCR)'
        name = testid + '_' + pwd
        if pcrtest == 'negative':
            pcrtest = '陰性(未檢出核酸) / Negative 【CT＞40】'
        elif pcrtest == 'positive':
            pcrtest = '陽性 / Positive'
        else:
            print("<h3>查無PCR結果<h3>")
            conn.close()
            sys.exit()
        run = table0.cell(12,0).paragraphs[0].add_run(pcrtest)
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255,0,0)
        table0.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table0.cell(12,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = table0.cell(8,1).paragraphs[0].add_run(testtype)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(8,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #報告日期
        rdate = rdat[:-3]
        run = table0.cell(8,3).paragraphs[0].add_run(rdate)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(8,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #檢測方法
        run = table0.cell(9,1).paragraphs[0].add_run(method)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(9,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #醫檢師簽名檔
        if vuser2 == 'H123160258':#鄭偉志
            pic = '/var/www/html/pdf_reports/pic/0550.png'
        elif vuser2 == 'P124237860':#陳奕勳
            pic = '/var/www/html/pdf_reports/pic/0559.png'  
        elif vuser2 == 'F228002368':#林庭萱
            pic = '/var/www/html/pdf_reports/pic/0633.png'
        elif vuser2 == 'E125023479':#張本樺
            pic = '/var/www/html/pdf_reports/pic/0637.png'
        elif vuser2 == 'F128633398':#朱奕翰
            pic = '/var/www/html/pdf_reports/pic/0648.png'   
        else:
            pic = ''
        
        if pic != '':
            run = table1.cell(0,2).paragraphs[0].add_run()
            run.add_picture(pic)
            table1.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            pass
        
        #醫檢師、實驗室主管簽名日期
        signdate = datetime.datetime.now().strftime('%Y-%m-%d')
        run = table1.cell(0,3).paragraphs[0].add_run(signdate)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        table1.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table1.cell(0,3).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        run = table1.cell(0,7).paragraphs[0].add_run(signdate)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        table1.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table1.cell(0,7).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        if tdate in os.listdir('/var/www/html/pdf_reports/'):
            pass
        else:
            os.system('mkdir /var/www/html/pdf_reports/' + tdate)
            os.system('chmod 777 /var/www/html/pdf_reports/' + tdate)
        if name + '.docx' in os.listdir('/var/www/html/pdf_reports/' + tdate + '/'):
            os.system('rm /var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
        else:
            pass
        if name + '.pdf' in os.listdir('/var/www/html/pdf_reports/' + tdate + '/'):
            os.system('rm /var/www/html/pdf_reports/' + tdate + '/' + name + '.pdf')
        else:
            pass
        if testid + '.pdf' in os.listdir('/var/www/html/pdf_reports/' + tdate + '/'):
            os.system('rm /var/www/html/pdf_reports/' + tdate + '/' + testid + '.pdf')
        else:
            pass
        doc.save('/var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
        os.system('chmod 777 /var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
        
        #path = '/var/www/html/pdf_reports/' + tdat.split(' ')[0] + '/'
        #times = 0
        #while times < 100:
        #    if testid + '.pdf' in os.listdir('/var/www/html/pdf_reports/' + tdate + '/'):
        #        cmd = "UPDATE covid_trans SET pcrpdfflag='Y' WHERE sampleid2='" + testid + "'"
        #        cursor.execute(cmd)
        #        conn.commit()
        #        print('<h3>' + name.split('_')[0] + ' 已產生報告</h3>')
        #        break
        #    else:
        #        time.sleep(10)
    else:
        #打開模板
        doc = docx.Document('/var/www/html/templates/covid_templates_v31.docx')
        doc.styles['Normal'].font.name = u'微軟正黑體'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')
        
        Tables = doc.tables
        
        #table
        table0 = Tables[0]
        table1 = Tables[1]
        
        #中英姓名
        ename = fname + ' ' + lname
        if cname.strip() != '' and ename.strip() != '':
            run = table0.cell(2,1).paragraphs[0].add_run(cname + ' / ' + ename.strip())
        elif ename.strip() == '':
            run = table0.cell(2,1).paragraphs[0].add_run(cname)
        elif cname.strip() == '':
            run = table0.cell(2,1).paragraphs[0].add_run(ename)
        else:
            pass
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(2,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #生日
        run = table0.cell(2,3).paragraphs[0].add_run(dob.replace('-', '.'))
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(2,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #性別
        run = table0.cell(3,1).paragraphs[0].add_run(sex)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(3,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #手機
        run = table0.cell(3,3).paragraphs[0].add_run(mobile)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(3,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #身分證或居留證
        if userid.strip() != '':
            idcardno = userid.strip()
            run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
            run.font.size = Pt(9)
            run.font.bold = True
            table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.font.bold = True
            table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        elif residentpermit.strip() == 'Y':
            idcardno = hicardno.strip()
            run = table0.cell(4,0).paragraphs[0].add_run("居留證證號")
            run.font.size = Pt(9)
            run.font.bold = True
            table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = table0.cell(4,0).paragraphs[1].add_run("(ARC No.)")
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.font.bold = True
            table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            idcardno = 'NA'
            run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
            run.font.size = Pt(9)
            run.font.bold = True
            table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.font.bold = True
            table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        run = table0.cell(4,1).paragraphs[0].add_run(idcardno)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(4,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #護照
        if passportid.strip() == '':
            passportid = 'NA'
        else:
            pass
        run = table0.cell(4,3).paragraphs[0].add_run(passportid)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(4,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #e-mail
        if uemail.strip() == '':
            uemail = 'NA'
        else:
            pass
        run = table0.cell(5,1).paragraphs[0].add_run(uemail)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(5,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        #鼻咽
        # if xmrpturgency == '1':
        #     run = table0.cell(5,3).paragraphs[0].add_run("鼻咽拭子(Nasopharyngeal Swab)")
        #     run.font.size = Pt(9)
        #     run.font.name = 'Times New Roman'
        #     run.font.bold = True
        #     table0.cell(5,3).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        #     table0.cell(5,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # else:
        #     run = table0.cell(5,3).paragraphs[0].add_run("咽喉拭子 (Oropharyngeal Swab)")
        #     run.font.size = Pt(9)
        #     run.font.name = 'Times New Roman'
        #     run.font.bold = True
        #     table0.cell(5,3).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        #     table0.cell(5,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # run = table0.cell(5,3).paragraphs[0].add_run("咽喉拭子 (Oropharyngeal Swab)")
        # run.font.size = Pt(9)
        # run.font.name = 'Times New Roman'
        # run.font.bold = True
        # table0.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # table0.cell(5,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        #採樣日期
        tdate = tdat.split(' ')[0]
        tdatime = tdat[:-3]
        run = table0.cell(6,1).paragraphs[0].add_run(tdatime)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(6,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #報告類型
        #if app_type.strip() == '1':
        #    app_type = '個人'
        #elif app_type.strip() == '2':
        #    cmd = "SELECT companyname from company where companyid='" + sendname + "'"
        #    cursor.execute(cmd)
        #    company = cursor.fetchone()
        #    if company == 'None' or company is None:
        #        app_type = '團體'
        #    else:
        #        app_type = company[0]
                
        #run = table0.cell(6,3).paragraphs[0].add_run(app_type)
        #run.font.size = Pt(9)
        #run.font.name = 'Times New Roman'
        #run.font.bold = True
        #table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        #table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #台胞證
        if mtpid.strip() == '':
            mtpid = 'NA'
        else:
            pass
        run = table0.cell(6,3).paragraphs[0].add_run(mtpid)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
 
        #檢測日期
        inspectime = (datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(seconds = round(int((datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.datetime.strptime(tdat, '%Y-%m-%d %H:%M:%S')).seconds) / 2, 0))).strftime('%Y-%m-%d %H:%M')
        run = table0.cell(7,1).paragraphs[0].add_run(inspectime)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(7,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #檢測項目、檢測類型、方法、結果、檔案名
        testtype = "新冠肺炎病毒抗原快篩檢測\n(SARS-CoV-2 Virus Ag Rapid Testing)"
        method = '免疫層析法 (Immunochromatographic Assay)'
        name = testid + '_' + pwd
        if ftest == 'negative':
            ftest = '陰性 / Negative'
        elif ftest == 'positive':
            ftest = '陽性 / Positive'
        else:
            print("<h3>查無快篩結果<h3>")
            conn.close()
            sys.exit()
        run = table0.cell(12,0).paragraphs[0].add_run(ftest)
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255,0,0)
        table0.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table0.cell(12,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = table0.cell(8,1).paragraphs[0].add_run(testtype)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(8,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #報告日期
        rdate = rdat[:-3]
        run = table0.cell(8,3).paragraphs[0].add_run(rdate)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(8,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #檢測方法
        run = table0.cell(9,1).paragraphs[0].add_run(method)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        table0.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table0.cell(9,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        #醫檢師簽名檔
        if vuser2 == 'H123160258':#鄭偉志
            pic = '/var/www/html/pdf_reports/pic/0550.png'
        elif vuser2 == 'P124237860':#陳奕勳
            pic = '/var/www/html/pdf_reports/pic/0559.png'  
        elif vuser2 == 'F228002368':#林庭萱
            pic = '/var/www/html/pdf_reports/pic/0633.png'
        elif vuser2 == 'E125023479':#張本樺
            pic = '/var/www/html/pdf_reports/pic/0637.png'
        elif vuser2 == 'F128633398':#朱奕翰
            pic = '/var/www/html/pdf_reports/pic/0648.png' 
        else:
            pic = ''
        
        if pic != '':
            run = table1.cell(0,2).paragraphs[0].add_run()
            run.add_picture(pic)
            table1.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            pass
        
        #醫檢師、實驗室主管簽名日期
        signdate = datetime.datetime.now().strftime('%Y-%m-%d')
        run = table1.cell(0,3).paragraphs[0].add_run(signdate)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        table1.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table1.cell(0,3).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        run = table1.cell(0,7).paragraphs[0].add_run(signdate)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        table1.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table1.cell(0,7).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        if tdate in os.listdir('/var/www/html/pdf_reports/'):
            pass
        else:
            os.system('mkdir /var/www/html/pdf_reports/' + tdate)
            os.system('chmod 777 /var/www/html/pdf_reports/' + tdate)
        if name + '.docx' in os.listdir('/var/www/html/pdf_reports/' + tdate + '/'):
            os.system('rm /var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
        else:
            pass
        if name + '.pdf' in os.listdir('/var/www/html/pdf_reports/' + tdate + '/'):
            os.system('rm /var/www/html/pdf_reports/' + tdate + '/' + name + '.pdf')
        else:
            pass
        if testid + '.pdf' in os.listdir('/var/www/html/pdf_reports/' + tdate + '/'):
            os.system('rm /var/www/html/pdf_reports/' + tdate + '/' + testid + '.pdf')
        else:
            pass
        doc.save('/var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
        os.system('chmod 777 /var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
        
path = '/var/www/html/pdf_reports/' + tdate + '/'
os.system('libreoffice --headless --convert-to pdf:writer_pdf_Export --outdir ' + path + ' /var/www/html/pdf_reports/' + tdate + '/' + name + '.docx 1>/dev/null 2>&1')

input_file = '/var/www/html/pdf_reports/' + tdate + '/' + name + '.pdf'
output_file = '/var/www/html/pdf_reports/' + tdate + '/' + name.split('_')[0] + '.pdf'
pswd = name.split('_')[1]
#print(input_file,output_file,pswd)
pdf_add_pswd(input_file, output_file, pswd)
#os.system('rm ' + input_file)

if testid + '.pdf' in os.listdir(path):
    if testid[0] == 'F':
        cmd = "UPDATE covid_trans SET fpdfflag='Y' WHERE sampleid1='" + testid + "'"
        cursor.execute(cmd)
        conn.commit()
        print('<script>alert("' + name.split('_')[0] + ' 已產生報告");</script>')
    elif testid[0] == 'Q' or (re.match(r"[0-9]",testid[0])!=None):
        cmd = "UPDATE covid_trans SET pcrpdfflag='Y' WHERE sampleid2='" + testid + "'"
        cursor.execute(cmd)
        conn.commit()
        print('<script>alert("' + testid + ' 已產生報告");</script>')
    else:
        print('<script>alert("' + testid + ' 不符合編碼格式");</script>')
    
else:
    print('<script>alert("' + testid + ' 未產生報告");</script>')
print('<script>alert("檔案位於後台伺服器 ' + path + ' 路徑下");</script>')
conn.close()