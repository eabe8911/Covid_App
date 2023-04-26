# -*- encoding:utf-8 -*-
import os
import sys
import pymysql
import subprocess
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import datetime
#from datetime import datetime
from PyPDF2 import PdfFileReader as pdfReader
from PyPDF2 import PdfFileWriter as pdfWriter
import time
import traceback
import re
from time import gmtime, strftime

############  add password to pdf  ############
def pdf_add_pswd(input_file, output_file, pswd):
    in_file = open(input_file, "rb")
    in_pdf = pdfReader(in_file)
    out_pdf = pdfWriter()
    p = in_pdf.getPage(0)
    out_pdf.addPage(p)
    out_pdf.encrypt(pswd)
    out_file = open(output_file, "wb")
    out_pdf.write(out_file)
    in_file.close()
    out_file.close()

path = '/var/www/html/pdf_reports/'
while True:
    #連線mysql
    db_settings = {
                    "host": "localhost",
                    "port": 3306,
                    "user": "libo_user",
                    "password": "xxx",
                    "db": "libodb",
                }
    try:
        # 建立Connection物件
        conn = pymysql.connect(**db_settings)
    except Exception as ex:
        ex_str=str(ex)
        time=str(strftime("%Y-%m-%d %H:%M:%S", gmtime()))
        print("current time:"+time+"\n"+ex_str+"\n")
        with open("autocreatereport_log.txt", "a",encoding="utf-8") as file_object:
            file_object.write("\nthis is ex message: \ncurrent time: "+time+"\n"+ex_str)
            file_object.close()

    cursor = conn.cursor()
    
    cmd = '''SELECT uuid,userid,mtpid,sex,cname,fname,lname,dob,mobile,uemail,ftest,pcrtest,type,telephone,address2,testtype,
            passportid,sampleid1,sampleid2,residentpermit,vuser2,tdat,rdat,sendname,frptflag,qrptflag,hicardno,fpdfflag,pcrpdfflag,mtpid
            FROM covid_trans WHERE 1=1 
            and ((lower(frptflag)='c' and fpdfflag='' and sampleid1 like 'F%') or (lower(qrptflag)='c' and pcrpdfflag='' and (sampleid2 like 'Q%' or sampleid2 like '0%' or sampleid2 like '1%' or sampleid2 like '2%' or sampleid2 like '3%' or sampleid2 like '4%' or sampleid2 like '5%' or sampleid2 like '6%' or sampleid2 like '7%' or sampleid2 like '8%' or sampleid2 like '9%')))
            '''
    cursor.execute(cmd)
    for Infos in cursor.fetchall():
        try:
            uuid,userid,mtpid,sex,cname,fname,lname,dob,mobile,uemail,ftest,pcrtest,app_type,telephone,address2,testtype,passportid,sampleid1,sampleid2,residentpermit,vuser2,tdat,rdat,sendname,frptflag,qrptflag,hicardno,fpdfflag,pcrpdfflag,mtpid = Infos[:]
            tdate = str(tdat).split(' ')[0]
            dob = str(dob)
            tdat = str(tdat)
            rdat = str(rdat)
            fname = fname.upper()
            lname = lname.upper()
            if userid.strip() != '' and userid.strip() != 'NA':
                pwd = userid.strip()
            elif residentpermit == 'Y':
                pwd = userid.strip()
            elif passportid.strip() != '' and passportid.strip() != 'NA':
                pwd = passportid.strip()
            else:
                pwd = dob.strip().replace('-', '').replace('/', '')
                
            if tdate in os.listdir(path):
                pass
            else:
                os.system('mkdir ' + path + tdate)
                os.system('chmod 777 ' + path + tdate)
            
            #產word
            if testtype == '1' or testtype == '2':
                #打開模板
                doc = docx.Document('/var/www/html/templates/covid_templates_v31.docx')
                doc.styles['Normal'].font.name = u'微軟正黑體'
                doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')
                
                Tables = doc.tables
                
                #table
                table0 = Tables[0]
                table1 = Tables[1]
                #table2 = Tables[2]
            
                #中英姓名
                ename = fname + ' ' + lname
                if cname.strip() != '' and ename.strip() != '':
                    run = table0.cell(2,1).paragraphs[0].add_run(cname + ' / ' + ename.strip())
                elif ename.strip() == '':
                    run = table0.cell(2,1).paragraphs[0].add_run(cname)
                elif cname.strip() == '':
                    run = table0.cell(2,1).paragraphs[0].add_run(ename)
                else:
                    pass
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(2,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #生日
                run = table0.cell(2,3).paragraphs[0].add_run(str(dob).replace('-', '.'))
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(2,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #性別
                run = table0.cell(3,1).paragraphs[0].add_run(sex)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(3,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #手機
                run = table0.cell(3,3).paragraphs[0].add_run(mobile)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(3,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #身分證或居留證
                if userid.strip() != '':
                    idcardno = userid.strip()
                    run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                    run.font.size = Pt(9)
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif residentpermit.strip() == 'Y':
                    idcardno = userid.strip()
                    run = table0.cell(4,0).paragraphs[0].add_run("居留證證號")
                    run.font.size = Pt(9)
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = table0.cell(4,0).paragraphs[1].add_run("(ARC No.)")
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                else:
                    idcardno = 'NA'
                    run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                    run.font.size = Pt(9)
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                run = table0.cell(4,1).paragraphs[0].add_run(idcardno)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(4,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #護照
                if passportid.strip() == '':
                    passportid = 'NA'
                else:
                    pass
                run = table0.cell(4,3).paragraphs[0].add_run(passportid)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(4,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #e-mail
                if uemail.strip() == '':
                    uemail = 'NA'
                else:
                    pass
                run = table0.cell(5,1).paragraphs[0].add_run(uemail)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(5,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                # #鼻咽
                # if throat.strip() == '鼻':
                #     throat = '鼻咽拭子(Nasopharyngeal Swab)'
                # else:
                #     throat = '咽喉拭子 (Oropharyngeal Swab)'
                # run = table0.cell(5,4).paragraphs[0].add_run(throat)
                # run.font.size = Pt(9)
                # run.font.name = 'Times New Roman'
                # run.font.bold = True
                # table0.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # table0.cell(5,4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                
                #採樣日期
                tdate = str(tdat).split(' ')[0]
                tdatime = str(tdat)[:-3]
                run = table0.cell(6,1).paragraphs[0].add_run(tdatime)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(6,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # #報告類型
                # if app_type.strip() == '1':
                #     app_type = '個人'
                # elif app_type.strip() == '2':
                #     cmd = "SELECT companyname from company where companyid='" + sendname + "'"
                #     cursor.execute(cmd)
                #     company = cursor.fetchone()
                #     if company == 'None' or company is None:
                #         app_type = '團體'
                #     else:
                #         app_type = company[0]
                        
                # run = table0.cell(6,3).paragraphs[0].add_run(app_type)
                # run.font.size = Pt(9)
                # run.font.name = 'Times New Roman'
                # run.font.bold = True
                # table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                #台胞證
                if mtpid.strip() == '':
                   mtpid = 'NA'
                else:
                   pass
                run = table0.cell(6,3).paragraphs[0].add_run(mtpid)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                
                #檢測項目
                if testtype == '1':
                    #檢測類型、方法、結果、檔案名
                    testtype = "新冠肺炎病毒抗原快篩檢測\n(SARS-CoV-2 Virus Ag Rapid Testing)"
                    method = '免疫層析法 (Immunochromatographic Assay)'
                    name = sampleid1 + '_' + pwd
                    if ftest == 'negative':
                        ftest = '陰性 / Negative'
                    elif ftest == 'positive':
                        ftest = '陽性 / Positive'
                    else:
                        continue
                    run = table0.cell(12,0).paragraphs[0].add_run(ftest)
                    run.font.size = Pt(16)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,0,0)
                    table0.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table0.cell(12,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    #檢測日期 ((報告時間+報到時間)/2)
                    inspectime = (datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(seconds = round(int((datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.datetime.strptime(tdat, '%Y-%m-%d %H:%M:%S')).seconds) / 2, 0))).strftime('%Y-%m-%d %H:%M')
                else:
                    #檢測類型、方法、結果、檔案名
                    testtype = "新冠肺炎病毒核酸檢測\n(SARS-CoV-2 Virus PCR Testing)"
                    method = '即時反轉錄聚合酶連鎖反應 (Real-time RT-PCR)'
                    name = sampleid2 + '_' + pwd
                    if pcrtest == 'negative':
                        pcrtest = '陰性 / Negative'
                    elif pcrtest == 'positive':
                        pcrtest = '陽性 / Positive'
                    else:
                        continue
                    run = table0.cell(12,0).paragraphs[0].add_run(pcrtest)
                    run.font.size = Pt(16)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,0,0)
                    table0.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table0.cell(12,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #檢測日期: 報告時間- 80 mins
                    inspectime = datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(minutes = 80)
                    if int(inspectime.strftime('%Y%m%d%H%M%S')) > int(tdat.replace('-', '').replace(':', '').replace(' ', '')):
                        inspectime = inspectime.strftime('%Y-%m-%d %H:%M')
                    else:
                        inspectime = (datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(seconds = round(int((datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.datetime.strptime(tdat, '%Y-%m-%d %H:%M:%S')).seconds) / 2, 0))).strftime('%Y-%m-%d %H:%M')
                    
                run = table0.cell(8,1).paragraphs[0].add_run(testtype)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(8,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #檢測日期
                run = table0.cell(7,1).paragraphs[0].add_run(inspectime)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(7,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #報告日期
                rdate = str(rdat)[:-3]
                run = table0.cell(8,3).paragraphs[0].add_run(rdate)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(8,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #檢測方法
                run = table0.cell(9,1).paragraphs[0].add_run(method)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(9,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #醫檢師簽名檔
                if vuser2 == 'B122408253':#陳柏源-離職
                    pic = '/var/www/html/pdf_reports/pic/0570.png'
                elif vuser2 == 'N123478768':#陳毓峻-離職
                    pic = '/var/www/html/pdf_reports/pic/0564.png'
                elif vuser2 == 'H123160258':#鄭偉志
                    pic = '/var/www/html/pdf_reports/pic/0550.png'
                elif vuser2 == 'P124237860':#陳奕勳
                    pic = '/var/www/html/pdf_reports/pic/0559.png'    
                elif vuser2 == 'P222717661':#李桂榕-離職
                    pic = '/var/www/html/pdf_reports/pic/0544.png'  
                elif vuser2 == 'N225198185':#黃琴涵-離職
                    pic = '/var/www/html/pdf_reports/pic/0561.png'  
                elif vuser2 == 'A225558000':#許育華-離職
                    pic = '/var/www/html/pdf_reports/pic/0513.png'  
                elif vuser2 == 'D221222459':#石鴻瑞
                    pic = '/var/www/html/pdf_reports/pic/0581.png'  
                elif vuser2 == 'R220853399':#曾彥瑜
                    pic = '/var/www/html/pdf_reports/pic/0558.png'
                else:
                    pic = ''
                
                if pic != '':
                    run = table1.cell(0,2).paragraphs[0].add_run()
                    run.add_picture(pic)
                    table1.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    pass
                
                #醫檢師、實驗室主管簽名日期
                signdate = datetime.datetime.now().strftime('%Y-%m-%d')
                run = table1.cell(0,3).paragraphs[0].add_run(signdate)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                table1.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table1.cell(0,3).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                run = table1.cell(0,7).paragraphs[0].add_run(signdate)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                table1.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table1.cell(0,7).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                doc.save('/var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
                #產pdf
                os.system('doc2pdf /var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
                input_file = '/var/www/html/pdf_reports/' + tdate + '/' + name + '.pdf'
                output_file = '/var/www/html/pdf_reports/' + tdate + '/' + name.split('_')[0] + '.pdf'
                pswd = name.split('_')[1]
                print(input_file,output_file,pswd)
                pdf_add_pswd(input_file, output_file, pswd)
                #os.system('rm ' + input_file)
                
                if output_file.split('/')[-1] in os.listdir('/var/www/html/pdf_reports/' + tdate + '/'):
                    testid = name.split('_')[0]
                    if testid[0] == 'Q' or (re.match(r"[0-9]",testid[0])!=None):
                        cmd = "UPDATE covid_trans SET pcrpdfflag='Y' WHERE sampleid2='" + testid + "'"
                        cursor.execute(cmd)
                        conn.commit()
                        os.system('python3 /var/www/html/belle/write_pdf_report_auto.py {}'.format(testid))
                    elif testid[0] == 'F':
                        cmd = "UPDATE covid_trans SET fpdfflag='Y' WHERE sampleid1='" + testid + "'"
                        cursor.execute(cmd)
                        conn.commit()
                    else:
                        pass
                else:
                    pass
                
            elif testtype == '3':
                if (ftest == 'negative' or ftest == 'positive') and frptflag.lower() == 'c':
                    #打開模板
                    doc = docx.Document('/var/www/html/templates/covid_templates_v31.docx')
                    doc.styles['Normal'].font.name = u'微軟正黑體'
                    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')
                    
                    Tables = doc.tables
                    
                    #table
                    table0 = Tables[0]
                    table1 = Tables[1]
                    #table2 = Tables[2]
                
                    #中英姓名
                    ename = fname + ' ' + lname
                    if cname.strip() != '' and ename.strip() != '':
                        run = table0.cell(2,1).paragraphs[0].add_run(cname + ' / ' + ename.strip())
                    elif ename.strip() == '':
                        run = table0.cell(2,1).paragraphs[0].add_run(cname)
                    elif cname.strip() == '':
                        run = table0.cell(2,1).paragraphs[0].add_run(ename)
                    else:
                        pass
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(2,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #生日
                    run = table0.cell(2,3).paragraphs[0].add_run(str(dob).replace('-', '.'))
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(2,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #性別
                    run = table0.cell(3,1).paragraphs[0].add_run(sex)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(3,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #手機
                    run = table0.cell(3,3).paragraphs[0].add_run(mobile)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(3,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #身分證或居留證
                    if userid.strip() != '':
                        idcardno = userid.strip()
                        run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    elif residentpermit.strip() == 'Y':
                        idcardno = userid.strip()
                        run = table0.cell(4,0).paragraphs[0].add_run("居留證證號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(ARC No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    else:
                        idcardno = 'NA'
                        run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    run = table0.cell(4,1).paragraphs[0].add_run(idcardno)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #護照
                    if passportid.strip() == '':
                        passportid = 'NA'
                    else:
                        pass
                    run = table0.cell(4,3).paragraphs[0].add_run(passportid)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #e-mail
                    if uemail.strip() == '':
                        uemail = 'NA'
                    else:
                        pass
                    run = table0.cell(5,1).paragraphs[0].add_run(uemail)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(5,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # #鼻咽
                    # if throat.strip() == '鼻':
                    #     throat = '鼻咽拭子(Nasopharyngeal Swab)'
                    # else:
                    #     throat = '咽喉拭子 (Oropharyngeal Swab)'
                    # run = table0.cell(5,4).paragraphs[0].add_run(throat)
                    # run.font.size = Pt(9)
                    # run.font.name = 'Times New Roman'
                    # run.font.bold = True
                    # table0.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # table0.cell(5,4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    #採樣日期
                    tdate = str(tdat).split(' ')[0]
                    tdatime = str(tdat)[:-3]
                    run = table0.cell(6,1).paragraphs[0].add_run(tdatime)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(6,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #報告類型
                    # if app_type.strip() == '1':
                    #     app_type = '個人'
                    # elif app_type.strip() == '2':
                    #     cmd = "SELECT companyname from company where companyid='" + sendname + "'"
                    #     cursor.execute(cmd)
                    #     company = cursor.fetchone()
                    #     if company == 'None' or company is None:
                    #         app_type = '團體'
                    #     else:
                    #         app_type = company[0]
                       
                    # run = table0.cell(6,3).paragraphs[0].add_run(app_type)
                    # run.font.size = Pt(9)
                    # run.font.name = 'Times New Roman'
                    # run.font.bold = True
                    # table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    #台胞證
                    if mtpid.strip() == '':
                        mtpid = 'NA'
                    else:
                        pass
                    run = table0.cell(6,3).paragraphs[0].add_run(mtpid)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


                    
                    #檢測日期 ((報告時間+報到時間)/2)
                    inspectime = (datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(seconds = round(int((datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.datetime.strptime(tdat, '%Y-%m-%d %H:%M:%S')).seconds) / 2, 0))).strftime('%Y-%m-%d %H:%M')
                    run = table0.cell(7,1).paragraphs[0].add_run(inspectime)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(7,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #檢測項目、檢測類型、方法、結果
                    testtype = "新冠肺炎病毒抗原快篩檢測\n(SARS-CoV-2 Virus Ag Rapid Testing)"
                    method = '免疫層析法 (Immunochromatographic Assay)'
                    name = sampleid1 + '_' + pwd
                    if ftest == 'negative':
                        ftest = '陰性 / Negative'
                    else:
                        ftest = '陽性 / Positive'
                    run = table0.cell(12,0).paragraphs[0].add_run(ftest)
                    run.font.size = Pt(16)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,0,0)
                    table0.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table0.cell(12,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    run = table0.cell(8,1).paragraphs[0].add_run(testtype)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(8,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #報告日期
                    rdate = str(rdat)[:-3]
                    run = table0.cell(8,3).paragraphs[0].add_run(rdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(8,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #檢測方法
                    run = table0.cell(9,1).paragraphs[0].add_run(method)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(9,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #醫檢師簽名檔
                    if vuser2 == 'B122408253':#陳柏源-離職
                        pic = '/var/www/html/pdf_reports/pic/0570.png'
                    elif vuser2 == 'N123478768':#陳毓峻-離職
                        pic = '/var/www/html/pdf_reports/pic/0564.png'
                    elif vuser2 == 'H123160258':#鄭偉志
                        pic = '/var/www/html/pdf_reports/pic/0550.png'
                    elif vuser2 == 'P124237860':#陳奕勳
                        pic = '/var/www/html/pdf_reports/pic/0559.png'  
                    elif vuser2 == 'P222717661':#李桂榕-離職
                        pic = '/var/www/html/pdf_reports/pic/0544.png'  
                    elif vuser2 == 'N225198185':#黃琴涵-離職
                        pic = '/var/www/html/pdf_reports/pic/0561.png'  
                    elif vuser2 == 'A225558000':#許育華-離職
                        pic = '/var/www/html/pdf_reports/pic/0513.png' 
                    elif vuser2 == 'D221222459':#石鴻瑞
                        pic = '/var/www/html/pdf_reports/pic/0581.png'
                    elif vuser2 == 'R220853399':#曾彥瑜
                         pic = '/var/www/html/pdf_reports/pic/0558.png'    
                    else:
                        pic = ''
                    
                    if pic != '':
                        run = table1.cell(0,2).paragraphs[0].add_run()
                        run.add_picture(pic)
                        table1.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        pass
                    
                    #醫檢師、實驗室主管簽名日期
                    signdate = datetime.datetime.now().strftime('%Y-%m-%d')
                    run = table1.cell(0,3).paragraphs[0].add_run(signdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    table1.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table1.cell(0,3).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                    run = table1.cell(0,7).paragraphs[0].add_run(signdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    table1.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table1.cell(0,7).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                    doc.save('/var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
                    #產pdf
                    os.system('doc2pdf /var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
                    input_file = '/var/www/html/pdf_reports/' + tdate + '/' + name + '.pdf'
                    output_file = '/var/www/html/pdf_reports/' + tdate + '/' + name.split('_')[0] + '.pdf'
                    pswd = name.split('_')[1]
                    print(input_file,output_file,pswd)
                    pdf_add_pswd(input_file, output_file, pswd)
                    #os.system('rm ' + input_file)
                    
                    if output_file.split('/')[-1] in os.listdir('/var/www/html/pdf_reports/' + tdate + '/'):
                        testid = name.split('_')[0]
                        if testid[0] == 'Q' or (re.match(r"[0-9]",testid[0])!=None):
                            cmd = "UPDATE covid_trans SET pcrpdfflag='Y' WHERE sampleid2='" + testid + "'"
                            cursor.execute(cmd)
                            conn.commit()
                            os.system('python3 /var/www/html/belle/write_pdf_report_auto.py {}'.format(testid))
                        elif testid[0] == 'F':
                            cmd = "UPDATE covid_trans SET fpdfflag='Y' WHERE sampleid1='" + testid + "'"
                            cursor.execute(cmd)
                            conn.commit()
                        else:
                            pass
                else:
                    pass
                
                if (pcrtest == 'negative' or ftest == 'positive') and qrptflag.lower() == 'c':
                    #打開模板
                    doc = docx.Document('/var/www/html/templates/covid_templates_v31.docx')
                    doc.styles['Normal'].font.name = u'微軟正黑體'
                    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')
                    
                    Tables = doc.tables
                    
                    #table
                    table0 = Tables[0]
                    table1 = Tables[1]
                    #table2 = Tables[2]
                
                    #中英姓名
                    ename = fname + ' ' + lname
                    if cname.strip() != '' and ename.strip() != '':
                        run = table0.cell(2,1).paragraphs[0].add_run(cname + ' / ' + ename.strip())
                    elif ename.strip() == '':
                        run = table0.cell(2,1).paragraphs[0].add_run(cname)
                    elif cname.strip() == '':
                        run = table0.cell(2,1).paragraphs[0].add_run(ename)
                    else:
                        pass
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(2,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #生日
                    run = table0.cell(2,3).paragraphs[0].add_run(str(dob).replace('-', '.'))
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(2,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #性別
                    run = table0.cell(3,1).paragraphs[0].add_run(sex)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(3,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #手機
                    run = table0.cell(3,3).paragraphs[0].add_run(mobile)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(3,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #身分證或居留證
                    if userid.strip() != '':
                        idcardno = userid.strip()
                        run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    elif residentpermit.strip() == 'Y':
                        idcardno = userid.strip()
                        run = table0.cell(4,0).paragraphs[0].add_run("居留證證號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(ARC No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    else:
                        idcardno = 'NA'
                        run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    run = table0.cell(4,1).paragraphs[0].add_run(idcardno)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #護照
                    if passportid.strip() == '':
                        passportid = 'NA'
                    else:
                        pass
                    run = table0.cell(4,3).paragraphs[0].add_run(passportid)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #e-mail
                    if uemail.strip() == '':
                        uemail = 'NA'
                    else:
                        pass
                    run = table0.cell(5,1).paragraphs[0].add_run(uemail)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(5,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # #鼻咽
                    # if throat.strip() == '鼻':
                    #     throat = '鼻咽拭子(Nasopharyngeal Swab)'
                    # else:
                    #     throat = '咽喉拭子 (Oropharyngeal Swab)'
                    # run = table0.cell(5,4).paragraphs[0].add_run(throat)
                    # run.font.size = Pt(9)
                    # run.font.name = 'Times New Roman'
                    # run.font.bold = True
                    # table0.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # table0.cell(5,4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    #採樣日期
                    tdate = str(tdat).split(' ')[0]
                    tdatime = str(tdat)[:-3]
                    run = table0.cell(6,1).paragraphs[0].add_run(tdatime)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(6,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #報告類型
                    # if app_type.strip() == '1':
                    #     app_type = '個人'
                    # elif app_type.strip() == '2':
                    #     cmd = "SELECT companyname from company where companyid='" + sendname + "'"
                    #     cursor.execute(cmd)
                    #     company = cursor.fetchone()
                    #     if company == 'None' or company is None:
                    #         app_type = '團體'
                    #     else:
                    #         app_type = company[0]
                    # run = table0.cell(6,3).paragraphs[0].add_run(app_type)
                    # run.font.size = Pt(9)
                    # run.font.name = 'Times New Roman'
                    # run.font.bold = True
                    # table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                #台胞證
                    if mtpid.strip() == '':
                        mtpid = 'NA'
                    else:
                        pass
                    run = table0.cell(6,3).paragraphs[0].add_run(mtpid)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER



                    #檢測日期: 報告時間- 80 mins
                    inspectime = datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(minutes = 80)
                    if int(inspectime.strftime('%Y%m%d%H%M%S')) > int(tdat.replace('-', '').replace(':', '').replace(' ', '')):
                        inspectime = inspectime.strftime('%Y-%m-%d %H:%M')
                    else:
                        inspectime = (datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(seconds = round(int((datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.datetime.strptime(tdat, '%Y-%m-%d %H:%M:%S')).seconds) / 2, 0))).strftime('%Y-%m-%d %H:%M')
                    run = table0.cell(7,1).paragraphs[0].add_run(inspectime)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(7,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #檢測項目、檢測類型、方法、結果
                    testtype = "新冠肺炎病毒核酸檢測\n(SARS-CoV-2 Virus PCR Testing)"
                    method = '即時反轉錄聚合酶連鎖反應 (Real-time RT-PCR)'
                    name = sampleid2 + '_' + pwd
                    if pcrtest == 'negative':
                        pcrtest = '陰性 / Negative'
                    else:
                        pcrtest = '陽性 / Positive'
                    
                    run = table0.cell(12,0).paragraphs[0].add_run(pcrtest)
                    run.font.size = Pt(16)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,0,0)
                    table0.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table0.cell(12,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    run = table0.cell(8,1).paragraphs[0].add_run(testtype)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(8,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #報告日期
                    rdate = rdat = str(rdat)[:-3]
                    run = table0.cell(8,3).paragraphs[0].add_run(rdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(8,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #檢測方法
                    run = table0.cell(9,1).paragraphs[0].add_run(method)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(9,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #醫檢師簽名檔
                    if vuser2 == 'B122408253': #陳柏源-離職
                        pic = '/var/www/html/pdf_reports/pic/0570.png' 
                    elif vuser2 == 'N123478768':#陳毓峻-離職
                        pic = '/var/www/html/pdf_reports/pic/0564.png'
                    elif vuser2 == 'H123160258':#鄭偉志
                        pic = '/var/www/html/pdf_reports/pic/0550.png'
                    elif vuser2 == 'P124237860':#陳奕勳
                        pic = '/var/www/html/pdf_reports/pic/0559.png'  
                    elif vuser2 == 'P222717661':#李桂榕-離職
                        pic = '/var/www/html/pdf_reports/pic/0544.png'  
                    elif vuser2 == 'N225198185':#黃琴涵-離職
                        pic = '/var/www/html/pdf_reports/pic/0561.png'  
                    elif vuser2 == 'A225558000':#許育華-離職
                        pic = '/var/www/html/pdf_reports/pic/0513.png'  
                    elif vuser2 == 'D221222459':#石鴻瑞
                        pic = '/var/www/html/pdf_reports/pic/0581.png'   
                    elif vuser2 == 'R220853399':#曾彥瑜
                        pic = '/var/www/html/pdf_reports/pic/0558.png'
                    else:
                        pic = ''
                    if pic != '':
                        run = table1.cell(0,2).paragraphs[0].add_run()
                        run.add_picture(pic)
                        table1.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        pass
                    
                    #醫檢師、實驗室主管簽名日期
                    signdate = datetime.datetime.now().strftime('%Y-%m-%d')
                    run = table1.cell(0,3).paragraphs[0].add_run(signdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    table1.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table1.cell(0,3).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                    run = table1.cell(0,7).paragraphs[0].add_run(signdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    table1.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table1.cell(0,7).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                    doc.save('/var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
                    
                    #產pdf
                    os.system('doc2pdf /var/www/html/pdf_reports/' + tdate + '/' + name + '.docx')
                    input_file = '/var/www/html/pdf_reports/' + tdate + '/' + name + '.pdf'
                    output_file = '/var/www/html/pdf_reports/' + tdate + '/' + name.split('_')[0] + '.pdf'
                    pswd = name.split('_')[1]
                    print(input_file,output_file,pswd)
                    pdf_add_pswd(input_file, output_file, pswd)
                    #os.system('rm ' + input_file)
                    
                    if output_file.split('/')[-1] in os.listdir('/var/www/html/pdf_reports/' + tdate + '/'):
                        testid = name.split('_')[0]
                        if testid[0] == 'Q' or (re.match(r"[0-9]",testid[0])!=None):
                            cmd = "UPDATE covid_trans SET pcrpdfflag='Y' WHERE sampleid2='" + testid + "'"
                            cursor.execute(cmd)
                            conn.commit()
                            os.system('python3 /var/www/html/belle/write_pdf_report_auto.py {}'.format(testid))
                        elif testid[0] == 'F':
                            cmd = "UPDATE covid_trans SET fpdfflag='Y' WHERE sampleid1='" + testid + "'"
                            cursor.execute(cmd)
                            conn.commit()
                        else:
                            pass
                else:
                    pass
            else:
                continue
            
            
        except:
            traceback_str=str(traceback.format_exc())
            time=str(strftime("%Y-%m-%d %H:%M:%S", gmtime()))
            print("current time:"+time+"\n"+traceback_str+"\n")
            with open("autocreatereport_log.txt", "a",encoding="utf-8") as file_object:
                file_object.write("\nthis is traceback message: \ncurrent time: "+time+"\n"+traceback_str)
                file_object.close()
            
        
    print('Sleep for 60s')
    conn.close()
    time.sleep(60)
