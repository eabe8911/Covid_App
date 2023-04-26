#-*- encoding=utf-8 -*-
import os
import sys
import subprocess
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import datetime
import pymysql
from PyPDF2 import PdfFileReader as pdfReader
from PyPDF2 import PdfFileWriter as pdfWriter
import time
import traceback

def pdf_add_pswd(input_file, output_file, pswd):
    in_file = open(input_file, "rb")
    in_pdf = pdfReader(in_file)
    out_pdf = pdfWriter()
    p = in_pdf.getPage(0)
    out_pdf.addPage(p)
    out_pdf.encrypt(pswd)
    out_file = open(output_file, "wb")
    out_pdf.write(out_file)
    in_file.close()
    out_file.close()

#連線mysql
db_settings = {
                "host": "localhost",
                "port": 3306,
                "user": "libo_user",
                "password": "xxx",
                "db": "libodb",
            }
try:
    # 建立Connection物件
    conn = pymysql.connect(**db_settings)
except Exception as ex:
    print(ex)
cursor = conn.cursor()

with open(sys.argv[1], 'r') as f:
    line1 = f.readline()
    line = f.readline()
    uuid,userid,sex,cname,fname,lname,dob,mobile,uemail,ftest,pcrtest,app_type,telephone,address2,testtype,passportid,sampleid1,sampleid2,residentpermit,vuser2,tdat,rdat,sendname,frptflag,qrptflag,hicardno,fpdfflag,pcrpdfflag,mtpid = line.split('\n')[0].split('\t')[:]
    if not tdat.split(' ')[0] in os.listdir('/var/www/html/pdf_reports/'):
        os.system('mkdir /var/www/html/pdf_reports/' + tdat.split(' ')[0])
        os.system('chmod 777 /var/www/html/pdf_reports/' + tdat.split(' ')[0])
    #os.system('chmod 777 /var/www/html/pdf_reports/*')
    
    while line.strip() != "":
        try:
            uuid,userid,sex,cname,fname,lname,dob,mobile,uemail,ftest,pcrtest,app_type,telephone,address2,testtype,passportid,sampleid1,sampleid2,residentpermit,vuser2,tdat,rdat,sendname,frptflag,qrptflag,hicardno,fpdfflag,pcrpdfflag,mtpid = line.split('\n')[0].split('\t')[:]
            if userid.strip() != '' and userid.strip() != 'NA':
                pwd = userid.strip()
            # elif residentpermit == 'Y':
            #     pwd = userid.strip()
            elif passportid.strip() != '' and passportid.strip() != 'NA':
                pwd = passportid.strip()
            else:
                pwd = dob.strip().replace('-', '').replace('/', '')
            if testtype == '1' or testtype == '2':
                #打開模板
                doc = docx.Document('/var/www/html/templates/covid_templates_v31.docx')
                doc.styles['Normal'].font.name = u'微軟正黑體'
                doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')
                
                Tables = doc.tables
                
                #table
                table0 = Tables[0]
                table1 = Tables[1]
                #table2 = Tables[2]
            
                #中英姓名
                ename = fname + ' ' + lname
                if cname.strip() != '' and ename.strip() != '':
                    run = table0.cell(2,1).paragraphs[0].add_run(cname + ' / ' + ename.strip())
                elif ename.strip() == '':
                    run = table0.cell(2,1).paragraphs[0].add_run(cname)
                elif cname.strip() == '':
                    run = table0.cell(2,1).paragraphs[0].add_run(ename)
                else:
                    pass
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(2,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #生日
                run = table0.cell(2,3).paragraphs[0].add_run(dob.replace('-', '.'))
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(2,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #性別
                run = table0.cell(3,1).paragraphs[0].add_run(sex)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(3,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #手機
                run = table0.cell(3,3).paragraphs[0].add_run(mobile)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(3,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #身分證或居留證
                if userid.strip() != '':
                    idcardno = userid.strip()
                    run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                    run.font.size = Pt(9)
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif residentpermit.strip() == 'Y':
                    idcardno = userid.strip()
                    run = table0.cell(4,0).paragraphs[0].add_run("居留證證號")
                    run.font.size = Pt(9)
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = table0.cell(4,0).paragraphs[1].add_run("(APC No.)")
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                else:
                    idcardno = 'NA'
                    run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                    run.font.size = Pt(9)
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                run = table0.cell(4,1).paragraphs[0].add_run(idcardno)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(4,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #護照
                if passportid.strip() == '':
                    passportid = 'NA'
                else:
                    pass
                run = table0.cell(4,3).paragraphs[0].add_run(passportid)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(4,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #e-mail
                if uemail.strip() == '':
                    uemail = 'NA'
                else:
                    pass
                run = table0.cell(5,1).paragraphs[0].add_run(uemail)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(5,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # #鼻咽
                # if throat.strip() == '鼻':
                #     throat = '鼻咽拭子(Nasopharyngeal Swab)'
                # else:
                #     throat = '咽喉拭子 (Oropharyngeal Swab)'
                # run = table0.cell(5,4).paragraphs[0].add_run(throat)
                # run.font.size = Pt(9)
                # run.font.name = 'Times New Roman'
                # run.font.bold = True
                # table0.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # table0.cell(5,4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                #採樣日期
                tdatime = tdat[:-3]
                run = table0.cell(6,1).paragraphs[0].add_run(tdatime)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(6,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #報告類型
                #if app_type.strip() == '1':
                #    app_type = '個人'
                #elif app_type.strip() == '2':
                #    cmd = "SELECT companyname from company where companyid='" + sendname + "'"
                #    cursor.execute(cmd)
                #    company = cursor.fetchone()
                #    if company == 'None' or company is None:
                #        app_type = '團體'
                #    else:
                #        app_type = company[0]
                #        
                #run = table0.cell(6,3).paragraphs[0].add_run(app_type)
                #run.font.size = Pt(9)
                #run.font.name = 'Times New Roman'
                #run.font.bold = True
                #table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                #table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #台胞證
                if mtpid.strip() == '':
                    mtpid = 'NA'
                else:
                    pass
                run = table0.cell(6,3).paragraphs[0].add_run(mtpid)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


                #報告日期
                rdate = rdat[:-3]
                run = table0.cell(8,3).paragraphs[0].add_run(rdate)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(8,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #檢測項目
                if testtype == '1':
                    #檢測類型、方法、結果、檔案名
                    testtype = "新冠肺炎病毒抗原快篩檢測\n(SARS-CoV-2 Virus Ag Rapid Testing)"
                    method = '免疫層析法 (Immunochromatographic Assay)'
                    name = sampleid1 + '_' + pwd
                    if ftest == 'negative':
                        ftest = '陰性 / Negative'
                    elif ftest == 'positive':
                        ftest = '陽性 / Positive'
                    else:
                        line = f.readline()
                        continue
                    run = table0.cell(12,0).paragraphs[0].add_run(ftest)
                    run.font.size = Pt(16)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,0,0)
                    table0.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table0.cell(12,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #檢測日期 ((報告時間+報到時間)/2)
                    inspectime = (datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(seconds = round(int((datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.datetime.strptime(tdat, '%Y-%m-%d %H:%M:%S')).seconds) / 2, 0))).strftime('%Y-%m-%d %H:%M')
                else:
                    #檢測類型、方法、結果、檔案名
                    testtype = "新冠肺炎病毒核酸檢測\n(SARS-CoV-2 Virus PCR Testing)"
                    method = '即時反轉錄聚合酶連鎖反應 (Real-time RT-PCR)'
                    name = sampleid2 + '_' + pwd
                    if pcrtest == 'negative':
                        pcrtest = '陰性(未檢出核酸) / Negative 【CT＞40】'
                    elif pcrtest == 'positive':
                        pcrtest = '陽性 / Positive'
                    else:
                        line = f.readline()
                        continue
                    run = table0.cell(12,0).paragraphs[0].add_run(pcrtest)
                    run.font.size = Pt(16)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,0,0)
                    table0.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table0.cell(12,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #檢測日期: 報告時間- 80 mins
                    inspectime = datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(minutes = 80)
                    if int(inspectime.strftime('%Y%m%d%H%M%S')) > int(tdat.replace('-', '').replace(':', '').replace(' ', '')):
                        inspectime = inspectime.strftime('%Y-%m-%d %H:%M')
                    else:
                        inspectime = (datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(seconds = round(int((datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.datetime.strptime(tdat, '%Y-%m-%d %H:%M:%S')).seconds) / 2, 0))).strftime('%Y-%m-%d %H:%M')
                    
                run = table0.cell(8,1).paragraphs[0].add_run(testtype)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(8,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #檢測日期
                run = table0.cell(7,1).paragraphs[0].add_run(inspectime)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(7,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #檢測方法
                run = table0.cell(9,1).paragraphs[0].add_run(method)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.bold = True
                table0.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table0.cell(9,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                #醫檢師簽名檔
                if vuser2 == 'H123160258':#鄭偉志
                    pic = '/var/www/html/pdf_reports/pic/0550.png'
                elif vuser2 == 'P124237860':#陳奕勳
                    pic = '/var/www/html/pdf_reports/pic/0559.png' 
                elif vuser2 == 'F228002368':#林庭萱
                    pic = '/var/www/html/pdf_reports/pic/0633.png'
                elif vuser2 == 'E125023479':#張本樺
                    pic = '/var/www/html/pdf_reports/pic/0637.png'
                else:
                    pic = ''
                
                if pic != '':
                    run = table1.cell(0,2).paragraphs[0].add_run()
                    run.add_picture(pic)
                    table1.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    pass
                
                #醫檢師、實驗室主管簽名日期
                signdate = datetime.datetime.now().strftime('%Y-%m-%d')
                run = table1.cell(0,3).paragraphs[0].add_run(signdate)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                table1.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table1.cell(0,3).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                run = table1.cell(0,7).paragraphs[0].add_run(signdate)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                table1.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                table1.cell(0,7).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                doc.save('/var/www/html/pdf_reports/' + name + '.docx')
                
            elif testtype == '3':
                if (ftest == 'negative' or ftest == 'positive') and frptflag.lower() == 'c':
                    #打開模板
                    doc = docx.Document('/var/www/html/templates/covid_templates_v31.docx')
                    doc.styles['Normal'].font.name = u'微軟正黑體'
                    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')
                    
                    Tables = doc.tables
                    
                    #table
                    table0 = Tables[0]
                    table1 = Tables[1]
                    #table2 = Tables[2]
                
                    #中英姓名
                    ename = fname + ' ' + lname
                    if cname.strip() != '' and ename.strip() != '':
                        run = table0.cell(2,1).paragraphs[0].add_run(cname + ' / ' + ename.strip())
                    elif ename.strip() == '':
                        run = table0.cell(2,1).paragraphs[0].add_run(cname)
                    elif cname.strip() == '':
                        run = table0.cell(2,1).paragraphs[0].add_run(ename)
                    else:
                        pass
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(2,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #生日
                    run = table0.cell(2,3).paragraphs[0].add_run(dob.replace('-', '.'))
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(2,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #性別
                    run = table0.cell(3,1).paragraphs[0].add_run(sex)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(3,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #手機
                    run = table0.cell(3,3).paragraphs[0].add_run(mobile)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(3,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #身分證或居留證
                    if userid.strip() != '':
                        idcardno = userid.strip()
                        run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    elif residentpermit.strip() == 'Y':
                        idcardno = userid.strip()
                        run = table0.cell(4,0).paragraphs[0].add_run("居留證證號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(APC No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    else:
                        idcardno = 'NA'
                        run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    run = table0.cell(4,1).paragraphs[0].add_run(idcardno)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #護照
                    if passportid.strip() == '':
                        passportid = 'NA'
                    else:
                        pass
                    run = table0.cell(4,3).paragraphs[0].add_run(passportid)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #e-mail
                    if uemail.strip() == '':
                        uemail = 'NA'
                    else:
                        pass
                    run = table0.cell(5,1).paragraphs[0].add_run(uemail)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(5,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # #鼻咽
                    # if throat.strip() == '鼻':
                    #     throat = '鼻咽拭子(Nasopharyngeal Swab)'
                    # else:
                    #     throat = '咽喉拭子 (Oropharyngeal Swab)'
                    # run = table0.cell(5,4).paragraphs[0].add_run(throat)
                    # run.font.size = Pt(9)
                    # run.font.name = 'Times New Roman'
                    # run.font.bold = True
                    # table0.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # table0.cell(5,4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    #採樣日期
                    tdatime = tdat[:-3]
                    run = table0.cell(6,1).paragraphs[0].add_run(tdatime)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(6,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #報告類型
                    #if app_type.strip() == '1':
                    #    app_type = '個人'
                    #elif app_type.strip() == '2':
                    #    cmd = "SELECT companyname from company where companyid='" + sendname + "'"
                    #    cursor.execute(cmd)
                    #    company = cursor.fetchone()
                    #    if company == 'None' or company is None:
                    #        app_type = '團體'
                    #    else:
                    #        app_type = company[0]
                        
                    #run = table0.cell(6,3).paragraphs[0].add_run(app_type)
                    #run.font.size = Pt(9)
                    #run.font.name = 'Times New Roman'
                    #run.font.bold = True
                    #table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    #table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    #台胞證
                    if mtpid.strip() == '':
                        mtpid= 'NA'
                    else:
                        pass
                    run = table0.cell(6,3).paragraphs[0].add_run(mtpid)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    #檢測日期
                    inspectime = (datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(seconds = round(int((datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.datetime.strptime(tdat, '%Y-%m-%d %H:%M:%S')).seconds) / 2, 0))).strftime('%Y-%m-%d %H:%M')
                    run = table0.cell(7,1).paragraphs[0].add_run(inspectime)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(7,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #檢測項目、檢測類型、方法、結果
                    testtype = "新冠肺炎病毒抗原快篩檢測\n(SARS-CoV-2 Virus Ag Rapid Testing)"
                    method = '免疫層析法 (Immunochromatographic Assay)'
                    name = sampleid1 + '_' + pwd
                    if ftest == 'negative':
                        ftest = '陰性 / Negative'
                    else:
                        ftest = '陽性 / Positive'
                    run = table0.cell(12,0).paragraphs[0].add_run(ftest)
                    run.font.size = Pt(16)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,0,0)
                    table0.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table0.cell(12,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    run = table0.cell(8,1).paragraphs[0].add_run(testtype)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(8,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #報告日期
                    rdate = rdat[:-3]
                    run = table0.cell(8,3).paragraphs[0].add_run(rdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(8,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #檢測方法
                    run = table0.cell(9,1).paragraphs[0].add_run(method)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(9,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #醫檢師簽名檔
                    if vuser2 == 'H123160258':#鄭偉志
                        pic = '/var/www/html/pdf_reports/pic/0550.png'
                    elif vuser2 == 'P124237860':#陳奕勳
                        pic = '/var/www/html/pdf_reports/pic/0559.png' 
                    elif vuser2 == 'F228002368':#林庭萱
                        pic = '/var/www/html/pdf_reports/pic/0633.png'
                    elif vuser2 == 'E125023479':#張本樺
                        pic = '/var/www/html/pdf_reports/pic/0637.png'
                    else:
                        pic = ''
                    
                    if pic != '':
                        run = table1.cell(0,2).paragraphs[0].add_run()
                        run.add_picture(pic)
                        table1.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        pass
                        
                    #醫檢師、實驗室主管簽名日期
                    signdate = datetime.datetime.now().strftime('%Y-%m-%d')
                    run = table1.cell(0,3).paragraphs[0].add_run(signdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    table1.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table1.cell(0,3).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                    run = table1.cell(0,7).paragraphs[0].add_run(signdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    table1.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table1.cell(0,7).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                    doc.save('/var/www/html/pdf_reports/' + name + '.docx')
                else:
                    pass
                
                if (pcrtest == 'negative' or ftest == 'positive') and qrptflag.lower() == 'c':
                    #打開模板
                    doc = docx.Document('/var/www/html/templates/covid_templates_v31.docx')
                    doc.styles['Normal'].font.name = u'微軟正黑體'
                    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')
                    
                    Tables = doc.tables
                    
                    #table
                    table0 = Tables[0]
                    table1 = Tables[1]
                    #table2 = Tables[2]
                
                    #中英姓名
                    ename = fname + ' ' + lname
                    if cname.strip() != '' and ename.strip() != '':
                        run = table0.cell(2,1).paragraphs[0].add_run(cname + ' / ' + ename.strip())
                    elif ename.strip() == '':
                        run = table0.cell(2,1).paragraphs[0].add_run(cname)
                    elif cname.strip() == '':
                        run = table0.cell(2,1).paragraphs[0].add_run(ename)
                    else:
                        pass
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(2,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #生日
                    run = table0.cell(2,3).paragraphs[0].add_run(dob.replace('-', '.'))
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(2,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #性別
                    run = table0.cell(3,1).paragraphs[0].add_run(sex)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(3,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #手機
                    run = table0.cell(3,3).paragraphs[0].add_run(mobile)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(3,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #身分證或居留證
                    if userid.strip() != '':
                        idcardno = userid.strip()
                        run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    elif residentpermit.strip() == 'Y':
                        idcardno = userid.strip()
                        run = table0.cell(4,0).paragraphs[0].add_run("居留證證號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(APC No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    else:
                        idcardno = 'NA'
                        run = table0.cell(4,0).paragraphs[0].add_run("身分證字號")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = table0.cell(4,0).paragraphs[1].add_run("(ID No.)")
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        table0.cell(4,0).paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        table0.cell(4,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    run = table0.cell(4,1).paragraphs[0].add_run(idcardno)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #護照
                    if passportid.strip() == '':
                        passportid = 'NA'
                    else:
                        pass
                    run = table0.cell(4,3).paragraphs[0].add_run(passportid)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(4,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #e-mail
                    if uemail.strip() == '':
                        uemail = 'NA'
                    else:
                        pass
                    run = table0.cell(5,1).paragraphs[0].add_run(uemail)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(5,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    # #鼻咽
                    # if throat.strip() == '鼻':
                    #     throat = '鼻咽拭子(Nasopharyngeal Swab)'
                    # else:
                    #     throat = '咽喉拭子 (Oropharyngeal Swab)'
                    # run = table0.cell(5,4).paragraphs[0].add_run(throat)
                    # run.font.size = Pt(9)
                    # run.font.name = 'Times New Roman'
                    # run.font.bold = True
                    # table0.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # table0.cell(5,4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER     

                    #採樣日期
                    tdatime = tdat[:-3]
                    run = table0.cell(6,1).paragraphs[0].add_run(tdatime)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(6,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #報告類型
                    #if app_type.strip() == '1':
                    #    app_type = '個人'
                    #elif app_type.strip() == '2':
                    #    cmd = "SELECT companyname from company where companyid='" + sendname + "'"
                    #    cursor.execute(cmd)
                    #    company = cursor.fetchone()
                    #    if company == 'None' or company is None:
                    #        app_type = '團體'
                    #    else:
                    #        app_type = company[0]
                    #run = table0.cell(6,3).paragraphs[0].add_run(app_type)
                    #run.font.size = Pt(9)
                    #run.font.name = 'Times New Roman'
                    #run.font.bold = True
                    #table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    #table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                   
                    #台胞證
                    if mtpid.strip() == '':
                        mtpid = 'NA'
                    else:
                        pass
                    run = table0.cell(6,3).paragraphs[0].add_run(mtpid)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(6,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    
                    #檢測日期: 報告時間 - 80 mins
                    inspectime = datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(minutes = 80)
                    if int(inspectime.strftime('%Y%m%d%H%M%S')) > int(tdat.replace('-', '').replace(':', '').replace(' ', '')):
                        inspectime = inspectime.strftime('%Y-%m-%d %H:%M')
                    else:
                        inspectime = (datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.timedelta(seconds = round(int((datetime.datetime.strptime(rdat, '%Y-%m-%d %H:%M:%S') - datetime.datetime.strptime(tdat, '%Y-%m-%d %H:%M:%S')).seconds) / 2, 0))).strftime('%Y-%m-%d %H:%M')
                    
                    run = table0.cell(7,1).paragraphs[0].add_run(inspectime)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(7,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #檢測項目、檢測類型、方法、結果
                    testtype = "新冠肺炎病毒核酸檢測\n(SARS-CoV-2 Virus PCR Testing)"
                    method = '即時反轉錄聚合酶連鎖反應 (Real-time RT-PCR)'
                    name = sampleid2 + '_' + pwd
                    if pcrtest == 'negative':
                        pcrtest = '陰性(未檢出核酸) / Negative 【CT＞40】'
                    else:
                        pcrtest = '陽性 / Positive'
                    
                    run = table0.cell(12,0).paragraphs[0].add_run(pcrtest)
                    run.font.size = Pt(16)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,0,0)
                    table0.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table0.cell(12,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    run = table0.cell(8,1).paragraphs[0].add_run(testtype)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(8,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #報告日期
                    rdate = rdat[:-3]
                    run = table0.cell(8,3).paragraphs[0].add_run(rdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(8,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #檢測方法
                    run = table0.cell(9,1).paragraphs[0].add_run(method)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    table0.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table0.cell(9,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    #醫檢師簽名檔
                    if vuser2 == 'H123160258':#鄭偉志
                        pic = '/var/www/html/pdf_reports/pic/0550.png'
                    elif vuser2 == 'P124237860':#陳奕勳
                        pic = '/var/www/html/pdf_reports/pic/0559.png' 
                    elif vuser2 == 'F228002368':#林庭萱
                        pic = '/var/www/html/pdf_reports/pic/0633.png'
                    elif vuser2 == 'E125023479':#張本樺
                        pic = '/var/www/html/pdf_reports/pic/0637.png'
                    else:
                        pic = ''
                    
                    if pic != '':
                        run = table1.cell(0,2).paragraphs[0].add_run()
                        run.add_picture(pic)
                        table1.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        pass
                    
                    #醫檢師、實驗室主管簽名日期
                    signdate = datetime.datetime.now().strftime('%Y-%m-%d')
                    run = table1.cell(0,3).paragraphs[0].add_run(signdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    table1.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table1.cell(0,3).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                    run = table1.cell(0,7).paragraphs[0].add_run(signdate)
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    table1.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    table1.cell(0,7).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                    doc.save('/var/www/html/pdf_reports/' + name + '.docx')
                else:
                    pass
            else:
                pass
            
            line = f.readline()
        except Exception as e:
            print(traceback.format_exc())
            line = f.readline()
            pass
#print(tdat.split(' ')[0])
os.system('cp /var/www/html/pdf_reports/*.docx /var/www/html/pdf_reports/' + tdat.split(' ')[0])
os.system('chmod 777 /var/www/html/pdf_reports/' + tdat.split(' ')[0] + '/*.docx')
os.system('rm /var/www/html/pdf_reports/*.docx')

#轉PDF、加密、取第一頁
path = '/var/www/html/pdf_reports/' + tdat.split(' ')[0] + '/'
os.system('libreoffice --headless --convert-to pdf:writer_pdf_Export --outdir ' + path + ' ' + path + '*.docx 1>/dev/null 2>&1')

for file in os.listdir(path):
    try:
        if file.split('.')[-1] == 'pdf' and '_' in file:
            input_file = path + file
            output_file = path + file.split('_')[0] + '.pdf'
            pswd = file.split('_')[1].split('.')[0]
            pdf_add_pswd(input_file, output_file, pswd)
            os.system('rm ' + input_file)
            ##更新資料庫
            testid = file.split('_')[0]
            if testid[0] == 'F':
                cmd = "UPDATE covid_trans SET fpdfflag='Y' WHERE sampleid1='" + testid + "'"
                cursor.execute(cmd)
                conn.commit()
            elif testid[0] == 'Q':
                cmd = "UPDATE covid_trans SET pcrpdfflag='Y' WHERE sampleid2='" + testid + "'"
                cursor.execute(cmd)
                conn.commit()
            else:
                continue
        else:
            continue
    except:
        errmsg = traceback.format_exc()
        print(errmsg.replace('\n', '<br>'))
    else:
        if testid + '.pdf' in os.listdir(path):
            print(testid + ' 已產生報告<br>')
        else:
            print(testid + ' 未產生報告<br>')
            
    
os.system('chmod 777 ' + path + '*.pdf')





#times = 0
#while times < 100:
#    #print(times)
#    m_docx = 0
#    m_pdf = 0
#    for i in os.listdir(path):
#        if '.docx' in i:
#            m_docx += 1
#        elif '.pdf' in i and not '_' in i:
#            m_pdf += 1
#    #print(m_docx)
#    #print(m_pdf)
#    try:
#        if m_docx == m_pdf:
#            for i in os.listdir(path):
#                try:
#                    if i.split('.')[-1] == 'pdf':
#                        testid = i.split('.')[0]
#                        if 'Q' == testid[0]:
#                            cmd = "UPDATE covid_trans set pcrpdfflag='Y' where sampleid2='" + testid + "' and lower(qrptflag)='c'"
#                        elif 'F' == testid[0]:
#                            cmd = "UPDATE covid_trans set fpdfflag='Y' where sampleid1='" + testid + "' and lower(frptflag)='c'"
#                        else:
#                            print(testid + ' 不符合編碼規則<br>')
#                            continue
#                        cursor.execute(cmd)
#                        conn.commit()
#                        print(testid + ' 已產生報告<br>')
#                    else:
#                        continue
#                except:
#                    print(testid + ' 未產生報告<br>')
#            break
#        else:
#            times += 1
#            time.sleep(10)
#    except Exception as e:
#        print(e)
    
print('<h4>報告已全部製作完成，放置於後台伺服器 ' + path + ' 路徑下</h4>')

conn.close()